from rest_framework.decorators import api_view
from rest_framework.response import Response
from docx import Document
from io import BytesIO
from ..models import DocumentoBoleto
from django.core.files.base import ContentFile
from datetime import datetime, timedelta
from rest_framework import viewsets
from ..serializers import DocumentoBoletoSerializer

class DocumentoBoletoViewSet(viewsets.ModelViewSet):
    queryset = DocumentoBoleto.objects.all()
    serializer_class = DocumentoBoletoSerializer


# Função para calcular a data de vencimento dos boletos
def calcular_data_vencimento(dias_uteis: int):
    data = datetime.today()
    while dias_uteis > 0:
        data += timedelta(days=1)
        if data.weekday() < 5:
            dias_uteis -= 1
    return data.strftime('%d/%m/%Y')


@api_view(['POST'])
def gerar_boleto_api(request):
    if request.method == "POST":
        dados = request.data
        # Carregar o template do documento
        document = Document('static/doc_exemplo.docx')
        placeholders = {
            'XNOMEX': dados.get('nome_cliente'),
            'XCPFX': dados.get('cpf_cnpj'),
            'XINSCRICAOX': dados.get('inscricao_estadual', ''),
            'XINSCRICAO_MUNICIPALX': dados.get('inscricao_municipal', ''),
            'XENDERECOX': dados.get('endereco'),
            'XCIDADEX': dados.get('cidade'),
            'XESTADOX': dados.get('estado'),
            'XCEPX': dados.get('cep'),
            'XEMAILX': dados.get('email'),
            'XFONEX': dados.get('telefone'),
            'XSERVICOX': dados.get('servico'),
            'XDATAX': calcular_data_vencimento(10),
        }

        # Substituir placeholders no documento
        for paragraph in document.paragraphs:
            for key, value in placeholders.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in placeholders.items():
                        if key in cell.text:
                            if key in ['XINSCRICAOX', 'XINSCRICAO_MUNICIPALX'] and not value:
                                value = ' '
                            cell.text = cell.text.replace(key, value)

        # Salvar o documento em memória
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        content = ContentFile(buffer.getvalue())

        # Salvar no banco de dados
        novo_documento = DocumentoBoleto(
            nome_cliente=dados.get('nome_cliente'),
        )
        novo_documento.arquivo.save(f'boleto_{dados["nome_cliente"]}.docx', content)

        # Retorna o link para o documento salvo
        return Response({"status": "Documento criado com sucesso!", "file_url": novo_documento.arquivo.url})
