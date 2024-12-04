from django.shortcuts import render, redirect
from django.http import HttpResponse
from docx import Document
from typing import Any
from .models import DocumentoBoleto
from django.core.files.base import ContentFile
from io import BytesIO
from datetime import datetime, timedelta

from .admin import CustomPagamentoForm

def home(request: Any) -> Any:
    return render(request, 'home.html')


def __calcular_data_vencimento(dias_uteis: int):
    data = datetime.today()
    while dias_uteis > 0:
        data += timedelta(days=1)
        if data.weekday() < 5:
            dias_uteis -= 1
    return data.strftime('%d/%m/%Y')


def gerar_boleto(request):
    if request.method == "POST":
        form = CustomPagamentoForm(request.POST)
        if form.is_valid():
            document = Document('static/doc_exemplo.docx')
            # Dados do formulário
            dados = {
                'XNOMEX': form.cleaned_data['nome_cliente'],
                'XCPFX': form.cleaned_data['cpf_cnpj'],
                'XINSCRICAOX': form.cleaned_data['inscricao_estadual'],
                'XINSCRICAO_MUNICIPALX': form.cleaned_data['inscricao_municipal'],
                'XENDERECOX': form.cleaned_data['endereco'],
                'XCIDADEX': form.cleaned_data['cidade'],
                'XESTADOX': form.cleaned_data['estado'],
                'XCEPX': form.cleaned_data['cep'],
                'XEMAILX': form.cleaned_data['email'],
                'XFONEX': form.cleaned_data['telefone'],
                'XSERVICOX': form.cleaned_data['servico'],
                'XDATAX': __calcular_data_vencimento(10),
            }

            # Substituir placeholders no documento
            for paragraph in document.paragraphs:
                for key, value in dados.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in dados.items():
                            if key in cell.text:
                                if key == 'XINSCRICAO_MUNICIPALX' or key == 'XINSCRICAOX':
                                    if not value:
                                        value = ' '
                                cell.text = cell.text.replace(key, value)

            # Salvar o documento em memória
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            content = ContentFile(buffer.getvalue())

            # Salvar no banco de dados
            novo_documento = DocumentoBoleto(
                nome_cliente=form.cleaned_data['nome_cliente'],
            )
            novo_documento.arquivo.save(f'boleto_{form.cleaned_data["nome_cliente"]}.docx', content)
            return redirect('lista_documentos')  # Redireciona para a página de documentos

    else:
        form = CustomPagamentoForm()
    return render(request, "pagamentos.html", {"form": form})



def lista_documentos(request):
    documentos = DocumentoBoleto.objects.all()
    return render(request, 'lista_documentos.html', {'documentos': documentos})

